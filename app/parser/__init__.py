import logging
import os

import docx
import fitz  # pymupdf

from typing import Optional

from app.models import ParsedDocument

logger = logging.getLogger(__name__)


def parse_file(filepath: str) -> Optional[ParsedDocument]:
    """根据文件扩展名解析文件内容"""
    ext = os.path.splitext(filepath)[1].lower()
    metadata = {"file_path": filepath, "extension": ext}

    try:
        if ext in (".md", ".txt"):
            text = _read_text(filepath)
        elif ext == ".pdf":
            text = _read_pdf(filepath)
        elif ext == ".docx":
            text = _read_docx(filepath)
        elif ext in (".java", ".xml", ".yaml", ".yml", ".sql"):
            text = _read_text(filepath)
            metadata["source_type"] = "code"
            metadata["language"] = ext.lstrip(".")
        else:
            return None

        if "source_type" not in metadata:
            metadata["source_type"] = "doc"

        return ParsedDocument(text=text, metadata=metadata)
    except Exception as e:
        logger.error(f"Failed to parse {filepath}: {e}")
        return None


def _read_text(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_pdf(filepath: str) -> str:
    doc = fitz.open(filepath)
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return text


def _read_docx(filepath: str) -> str:
    doc = docx.Document(filepath)
    return "\n".join(p.text for p in doc.paragraphs)
